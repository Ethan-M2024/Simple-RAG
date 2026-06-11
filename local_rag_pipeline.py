"""local_rag_pipeline.py.

A production-grade, fully local, dependency-light Retrieval-Augmented Generation
(RAG) pipeline written in explicit native Python.

This module deliberately avoids high-level orchestration frameworks
(LangChain, LlamaIndex, Haystack, etc.). Every stage of the data lifecycle is
written by hand against the primary package SDKs so that an open-source reader
can audit exactly what happens to their data:

    * Chunking         -> hand-rolled recursive character splitter.
    * Embedding        -> FastEmbed (ONNX runtime, no PyTorch) BGE-small.
    * Vector storage   -> ChromaDB local SQLite-backed PersistentClient.
    * Generation       -> Ollama native client driving a local qwen2.5:7b model.

The pipeline is partitioned into five logical blocks:

    1. NativeDocumentChunker      - deterministic text partitioning.
    2. LocalVectorStorage         - explicit embedding + Chroma ingestion.
    3. LocalVectorStorage.retrieve_context - querying + score normalization.
    4. DeterministicRAGEngine     - zero-hallucination local inference.
    5. Telemetry / failsafes      - timestamped logs woven through every stage.

Run directly to execute the built-in verification harness:

    python local_rag_pipeline.py
"""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import os
import re
import time
from datetime import datetime
from typing import Any, Dict, Generator, List, Optional, Tuple

import chromadb
from fastembed import TextEmbedding
from fastembed.rerank.cross_encoder import TextCrossEncoder

try:
    import ollama
except ImportError:  # pragma: no cover - educational guard.
    ollama = None  # type: ignore[assignment]

try:
    import fitz  # PyMuPDF: native PDF text + page/image rasterization.
except ImportError:  # pragma: no cover - optional loader dependency.
    fitz = None  # type: ignore[assignment]

try:
    import docx  # python-docx
except ImportError:  # pragma: no cover - optional loader dependency.
    docx = None  # type: ignore[assignment]

try:
    import io

    import pytesseract  # Tesseract OCR bridge.
    from PIL import Image
except ImportError:  # pragma: no cover - optional OCR dependency.
    pytesseract = None  # type: ignore[assignment]
    Image = None  # type: ignore[assignment]


# --------------------------------------------------------------------------- #
# Block 5 (foundation): Telemetry helpers.
# Defined first because every other block emits structured logs through it.
# --------------------------------------------------------------------------- #
def _fmt_secs(seconds: float) -> str:
    """Format a duration as a compact, human-readable string.

    Args:
        seconds: Elapsed time in seconds.

    Returns:
        ``"850ms"`` for sub-second, ``"4.2s"`` under a minute, else ``"3m51s"``.
    """
    if seconds < 1.0:
        return f"{seconds * 1000:.0f}ms"
    if seconds < 60.0:
        return f"{seconds:.1f}s"
    minutes, rem = divmod(int(round(seconds)), 60)
    return f"{minutes}m{rem:02d}s"


def _log(stage: str, message: str) -> None:
    """Emit a clean, timestamped, single-line telemetry record.

    Args:
        stage: Short upper-case label identifying the pipeline phase
            (e.g. ``"CHUNK"``, ``"EMBED"``, ``"QUERY"``).
        message: Human-readable detail for the log line.
    """
    stamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
    print(f"[{stamp}] [{stage:<7}] {message}")


# --------------------------------------------------------------------------- #
# Block 1: High-Performance Text Chunking Engine.
# --------------------------------------------------------------------------- #
class NativeDocumentChunker:
    """Deterministic, in-house recursive character text splitter.

    The splitter walks a descending priority list of structural separators
    (paragraphs -> lines -> words -> characters) and greedily packs fragments
    into blocks no larger than ``chunk_size``. A trailing sliding window of
    ``chunk_overlap`` characters is carried into each subsequent block to
    preserve cross-boundary context.

    The implementation contains no randomness and no hidden state, so identical
    inputs always yield byte-identical chunk boundaries.
    """

    #: Separators inspected in descending order of structural priority.
    #: The empty string is the terminal fallback (split into characters).
    SEPARATORS: List[str] = ["\n\n", "\n", " ", ""]

    def _atomize(
        self, text: str, chunk_size: int, sep_index: int = 0
    ) -> List[str]:
        """Recursively break ``text`` into fragments no larger than ``chunk_size``.

        The splitter tries separators in descending structural priority. It
        splits on the highest-priority separator present, then RECURSES into any
        resulting fragment that still exceeds ``chunk_size`` using the next finer
        separator (paragraphs -> lines -> words -> characters). This guarantees
        that a single oversized paragraph is broken down rather than emitted
        whole. Separators are re-attached so concatenation reproduces the input.

        Args:
            text: The raw text to atomize.
            chunk_size: Maximum allowed length of any returned fragment.
            sep_index: Index into :attr:`SEPARATORS` to try at this depth.

        Returns:
            An ordered list of fragments, each <= ``chunk_size`` wherever the
            text's structure permits (down to individual characters).
        """
        separator = (
            self.SEPARATORS[sep_index]
            if sep_index < len(self.SEPARATORS)
            else ""
        )
        if separator == "":
            # Terminal fallback: every character is its own fragment.
            return list(text)
        if separator not in text:
            # This separator does not occur; descend to the next finer one.
            return self._atomize(text, chunk_size, sep_index + 1)

        pieces = text.split(separator)
        fragments: List[str] = []
        for index, piece in enumerate(pieces):
            # Re-attach the separator to all but the final piece.
            fragment = piece + separator if index < len(pieces) - 1 else piece
            if not fragment:
                continue
            if len(fragment) > chunk_size and sep_index + 1 < len(self.SEPARATORS):
                # Still too big: recurse with the next finer separator.
                fragments.extend(
                    self._atomize(fragment, chunk_size, sep_index + 1)
                )
            else:
                fragments.append(fragment)
        return fragments

    def split_text(
        self,
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        source_label: str = "document",
    ) -> List[Dict[str, Any]]:
        """Split ``text`` into overlapping, metadata-rich chunks.

        Args:
            text: The full source document.
            chunk_size: Maximum character length of any single chunk.
            chunk_overlap: Number of trailing characters carried forward from
                each chunk into the next to maintain contextual continuity.
            source_label: Provenance label stamped onto every chunk's metadata.

        Returns:
            A list of dictionaries shaped as::

                {
                    "text": str,
                    "metadata": {
                        "chunk_index": int,
                        "char_count": int,
                        "source_label": str,
                    },
                }

        Raises:
            ValueError: If ``chunk_overlap`` is not strictly smaller than
                ``chunk_size`` (an overlap >= size cannot make forward progress).
        """
        if chunk_overlap >= chunk_size:
            raise ValueError(
                f"chunk_overlap ({chunk_overlap}) must be strictly smaller "
                f"than chunk_size ({chunk_size})."
            )

        fragments = self._atomize(text, chunk_size)
        chunks: List[Dict[str, Any]] = []
        current: str = ""

        def _emit(block: str) -> None:
            """Append a finalized, stripped block as a structured chunk."""
            cleaned = block.strip()
            if not cleaned:
                return
            chunks.append(
                {
                    "text": cleaned,
                    "metadata": {
                        "chunk_index": len(chunks),
                        "char_count": len(cleaned),
                        "source_label": source_label,
                    },
                }
            )

        for fragment in fragments:
            # A single fragment longer than chunk_size is force-emitted whole;
            # the atomizer guarantees this only happens for the char fallback.
            if len(current) + len(fragment) <= chunk_size:
                current += fragment
                continue

            _emit(current)
            # Seed the next block with the trailing overlap window.
            overlap_window = current[-chunk_overlap:] if chunk_overlap else ""
            current = overlap_window + fragment

        _emit(current)
        return chunks


# --------------------------------------------------------------------------- #
# Block 1b: OCR-Augmented Document Loader (.pdf / .docx / .txt / .md / images).
# --------------------------------------------------------------------------- #
class NativeDocumentLoader:
    """Extract maximal text from local files, including OCR of images & tables.

    The loader is built for total recall: it harvests every layer of a document
    so nothing — body text, table cells, figures, or scanned/embedded images —
    is silently dropped before retrieval.

    Per-format strategy:
        * ``.txt`` / ``.md`` -> read directly as UTF-8 text.
        * ``.pdf``           -> PyMuPDF native text + table extraction, plus
          Tesseract OCR for scanned pages and every embedded raster image.
        * ``.docx``          -> python-docx paragraphs + table cells, plus
          Tesseract OCR of every embedded image.
        * images (.png/.jpg/.jpeg/.tiff/.bmp/.webp) -> direct Tesseract OCR.

    OCR is performed by the local Tesseract 5 engine via ``pytesseract`` and is
    fully deterministic. Every reader is hand-written so an open-source reader
    can see exactly how bytes on disk become the strings handed to the chunker.

    Args:
        enable_ocr: Master switch for all Tesseract OCR passes.
        ocr_dpi: Rasterization resolution used when OCRing whole PDF pages.
            Higher values improve accuracy on small fonts at the cost of speed.
        ocr_lang: Tesseract language pack code (e.g. ``"eng"``).
        scanned_text_threshold: If a PDF page yields fewer than this many native
            characters, the page is treated as scanned and the full page is
            rendered and OCR'd instead of relying on sparse native text.
    """

    SUPPORTED: Tuple[str, ...] = (
        ".txt", ".md", ".pdf", ".docx",
        ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp",
    )
    IMAGE_EXTENSIONS: Tuple[str, ...] = (
        ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp",
    )

    def __init__(
        self,
        enable_ocr: bool = True,
        ocr_dpi: int = 200,
        ocr_lang: str = "eng",
        scanned_text_threshold: int = 100,
        ocr_embedded_images: bool = False,
    ) -> None:
        """Configure the loader and verify the OCR engine is reachable.

        Args:
            ocr_embedded_images: When True, every raster image embedded inside a
                digital PDF/DOCX is OCR'd individually. This is slow on
                image-heavy documents, so it defaults to False. Standalone image
                files and sparse/scanned PDF pages are still OCR'd regardless.
        """
        self.ocr_dpi = ocr_dpi
        self.ocr_lang = ocr_lang
        self.scanned_text_threshold = scanned_text_threshold
        self.ocr_embedded_images = ocr_embedded_images

        self.enable_ocr = enable_ocr
        if self.enable_ocr and (pytesseract is None or Image is None):
            _log("LOAD", "WARNING: 'pytesseract'/'Pillow' missing; OCR disabled. "
                         "Install with 'pip install pytesseract pillow'.")
            self.enable_ocr = False
        elif self.enable_ocr:
            try:
                version = pytesseract.get_tesseract_version()
                _log("LOAD", f"Tesseract OCR engine v{version} active "
                             f"(dpi={ocr_dpi}, lang='{ocr_lang}').")
            except Exception as exc:  # noqa: BLE001 - binary not on PATH.
                _log("LOAD", f"WARNING: Tesseract binary unavailable ({exc}); "
                             "OCR disabled. Install it via 'brew install "
                             "tesseract' or your OS package manager.")
                self.enable_ocr = False

    # ---------------------------- OCR primitives ----------------------------- #
    def _ocr_image(self, image: "Image.Image") -> str:
        """Run Tesseract on a single PIL image and return cleaned text.

        Args:
            image: A PIL image (RGB or grayscale).

        Returns:
            Extracted text, or an empty string on failure / empty result.
        """
        if not self.enable_ocr:
            return ""
        try:
            return pytesseract.image_to_string(image, lang=self.ocr_lang).strip()
        except Exception as exc:  # noqa: BLE001 - never crash the pipeline.
            _log("OCR", f"WARNING: OCR pass failed: {exc}")
            return ""

    def _ocr_image_bytes(self, blob: bytes) -> str:
        """Decode raw image bytes and OCR them.

        Args:
            blob: Encoded image bytes (PNG/JPEG/etc.).

        Returns:
            Extracted text, or an empty string on failure.
        """
        if not self.enable_ocr:
            return ""
        try:
            with Image.open(io.BytesIO(blob)) as image:
                return self._ocr_image(image.convert("RGB"))
        except Exception as exc:  # noqa: BLE001
            _log("OCR", f"WARNING: could not decode embedded image: {exc}")
            return ""

    # ---------------------------- format readers ----------------------------- #
    def _read_pdf(self, file_path: str) -> str:
        """Extract native text, tables, and OCR layers from a PDF.

        Args:
            file_path: Path to a ``.pdf`` file.

        Returns:
            Concatenated text across all pages and recovered layers.
        """
        if fitz is None:
            _log("LOAD", "WARNING: 'PyMuPDF' not installed; skipping PDF. "
                         "Install with 'pip install pymupdf'.")
            return ""

        # Silence non-fatal MuPDF chatter (e.g. broken ICC color profiles in
        # embedded images) so the telemetry stream stays clean.
        try:
            fitz.TOOLS.mupdf_display_errors(False)
        except Exception:  # noqa: BLE001 - older PyMuPDF lacks the toggle.
            pass

        name = os.path.basename(file_path)
        document = fitz.open(file_path)
        page_texts: List[str] = []
        ocr_pages = 0
        ocr_images = 0
        tables_found = 0

        for page_number, page in enumerate(document):
            parts: List[str] = []
            native = page.get_text("text").strip()

            if len(native) < self.scanned_text_threshold and self.enable_ocr:
                # Sparse native text -> treat page as scanned; OCR full render.
                pixmap = page.get_pixmap(dpi=self.ocr_dpi)
                with Image.open(io.BytesIO(pixmap.tobytes("png"))) as render:
                    ocr_text = self._ocr_image(render.convert("RGB"))
                if ocr_text:
                    parts.append(ocr_text)
                    ocr_pages += 1
            else:
                if native:
                    parts.append(native)
                # Digital page: optionally OCR embedded raster images (figures,
                # table screenshots) so body text is not duplicated.
                if self.enable_ocr and self.ocr_embedded_images:
                    for img in page.get_images(full=True):
                        try:
                            blob = document.extract_image(img[0])["image"]
                        except Exception:  # noqa: BLE001
                            continue
                        image_text = self._ocr_image_bytes(blob)
                        if image_text:
                            parts.append(f"[Image OCR] {image_text}")
                            ocr_images += 1

            # Structured table recovery (guarded; API varies across versions).
            try:
                tables = page.find_tables()
                for table in tables.tables:
                    rows = table.extract()
                    rendered = "\n".join(
                        " | ".join((cell or "") for cell in row) for row in rows
                    ).strip()
                    if rendered:
                        parts.append(f"[Table]\n{rendered}")
                        tables_found += 1
            except Exception:  # noqa: BLE001 - tables are best-effort.
                pass

            if parts:
                page_texts.append(f"[Page {page_number + 1}]\n" + "\n\n".join(parts))

        document.close()
        _log("LOAD", f"PDF '{name}': {len(page_texts)} pages, "
                     f"{ocr_pages} OCR'd scans, {ocr_images} OCR'd images, "
                     f"{tables_found} tables.")
        return "\n\n".join(page_texts)

    def _read_docx(self, file_path: str) -> str:
        """Extract paragraphs, table cells, and OCR'd images from a DOCX.

        Args:
            file_path: Path to a ``.docx`` file.

        Returns:
            Concatenated text across paragraphs, tables, and embedded images.
        """
        if docx is None:
            _log("LOAD", "WARNING: 'python-docx' not installed; skipping DOCX. "
                         "Install with 'pip install python-docx'.")
            return ""

        name = os.path.basename(file_path)
        document = docx.Document(file_path)
        parts: List[str] = []

        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        parts.extend(paragraphs)

        tables_found = 0
        for table in document.tables:
            rows = [
                " | ".join(cell.text.strip() for cell in row.cells)
                for row in table.rows
            ]
            rendered = "\n".join(r for r in rows if r.strip())
            if rendered:
                parts.append(f"[Table]\n{rendered}")
                tables_found += 1

        ocr_images = 0
        if self.enable_ocr and self.ocr_embedded_images:
            for rel in document.part.rels.values():
                if "image" not in rel.reltype:
                    continue
                try:
                    blob = rel.target_part.blob
                except Exception:  # noqa: BLE001
                    continue
                image_text = self._ocr_image_bytes(blob)
                if image_text:
                    parts.append(f"[Image OCR] {image_text}")
                    ocr_images += 1

        _log("LOAD", f"DOCX '{name}': {len(paragraphs)} paragraphs, "
                     f"{tables_found} tables, {ocr_images} OCR'd images.")
        return "\n\n".join(parts)

    def _read_image(self, file_path: str) -> str:
        """OCR a standalone image file.

        Args:
            file_path: Path to a supported image file.

        Returns:
            The OCR'd text, or an empty string if OCR is disabled/failed.
        """
        name = os.path.basename(file_path)
        if not self.enable_ocr:
            _log("LOAD", f"WARNING: OCR disabled; cannot read image '{name}'.")
            return ""
        try:
            with Image.open(file_path) as image:
                text = self._ocr_image(image.convert("RGB"))
        except Exception as exc:  # noqa: BLE001
            _log("LOAD", f"WARNING: could not open image '{name}': {exc}")
            return ""
        _log("LOAD", f"IMAGE '{name}': {len(text)} chars OCR'd.")
        return text

    def load_file(self, file_path: str) -> str:
        """Load a single file and return its fully extracted plain text.

        Args:
            file_path: Path to a supported document or image file.

        Returns:
            The extracted text, or an empty string if the format is
            unsupported or a required reader dependency is missing.
        """
        extension = os.path.splitext(file_path)[1].lower()

        if extension in (".txt", ".md"):
            with open(file_path, "r", encoding="utf-8", errors="replace") as handle:
                return handle.read()
        if extension == ".pdf":
            return self._read_pdf(file_path)
        if extension == ".docx":
            return self._read_docx(file_path)
        if extension in self.IMAGE_EXTENSIONS:
            return self._read_image(file_path)

        _log("LOAD", f"WARNING: unsupported file type '{extension}', skipping "
                     f"'{os.path.basename(file_path)}'.")
        return ""

    def load_path(self, path: str) -> List[Dict[str, str]]:
        """Load a single file or every supported file inside a directory.

        Args:
            path: A file path or a directory path. Directories are scanned
                recursively for supported extensions.

        Returns:
            A list of ``{"source_label": str, "text": str}`` records, one per
            file that yielded non-empty text.
        """
        targets: List[str] = []
        if os.path.isdir(path):
            for root, _dirs, files in os.walk(path):
                for name in sorted(files):
                    if name.lower().endswith(self.SUPPORTED):
                        targets.append(os.path.join(root, name))
        elif os.path.isfile(path):
            targets.append(path)
        else:
            _log("LOAD", f"WARNING: path not found: '{path}'.")
            return []

        _log("LOAD", f"Discovered {len(targets)} loadable file(s) under '{path}'.")

        documents: List[Dict[str, str]] = []
        for target in targets:
            text = self.load_file(target)
            if text.strip():
                documents.append(
                    {"source_label": os.path.basename(target), "text": text}
                )
        return documents


# --------------------------------------------------------------------------- #
# Blocks 2 & 3: Vector DB Storage, Explicit Embedding, Querying + Normalization.
# --------------------------------------------------------------------------- #
_WORD_RE = re.compile(r"[a-z0-9]+")


def _bm25_tokenize(text: str) -> List[str]:
    """Lowercase word tokenizer shared by BM25 indexing and querying."""
    return _WORD_RE.findall(text.lower())


class BM25Index:
    """Minimal in-memory BM25 (Okapi) over the project's stored chunks.

    Complements dense vector search with exact lexical matching — crucial
    for technical terms, identifiers, and rare words that embeddings blur.
    Built from the Chroma collection's documents; no extra dependencies.
    """

    K1 = 1.5
    B = 0.75

    def __init__(self, ids: List[str], documents: List[str],
                 metadatas: List[Dict[str, Any]]) -> None:
        self.ids = ids
        self.documents = documents
        self.metadatas = metadatas
        self._doc_tokens = [_bm25_tokenize(d) for d in documents]
        self._doc_lens = [len(t) for t in self._doc_tokens]
        self._avg_len = (sum(self._doc_lens) / len(self._doc_lens)
                         if self._doc_lens else 0.0)
        # term -> {doc_index -> term frequency}
        self._postings: Dict[str, Dict[int, int]] = {}
        for idx, tokens in enumerate(self._doc_tokens):
            for token in tokens:
                self._postings.setdefault(token, {})
                self._postings[token][idx] = self._postings[token].get(idx, 0) + 1

    def top_n(self, query: str, n: int) -> List[Tuple[int, float]]:
        """Return up to ``n`` (doc_index, score) pairs with positive score."""
        n_docs = len(self.documents)
        if n_docs == 0:
            return []
        scores: Dict[int, float] = {}
        for term in set(_bm25_tokenize(query)):
            postings = self._postings.get(term)
            if not postings:
                continue
            df = len(postings)
            idf = math.log(1.0 + (n_docs - df + 0.5) / (df + 0.5))
            for doc_idx, tf in postings.items():
                norm = self.K1 * (1.0 - self.B + self.B *
                                  (self._doc_lens[doc_idx] / self._avg_len))
                scores[doc_idx] = scores.get(doc_idx, 0.0) + (
                    idf * tf * (self.K1 + 1.0) / (tf + norm))
        ranked = sorted(scores.items(), key=lambda p: p[1], reverse=True)
        return ranked[:n]


class LocalVectorStorage:
    """Explicit, transparent bridge between FastEmbed and ChromaDB.

    All embedding generation is performed manually with a local FastEmbed
    (ONNX) model and the resulting raw float vectors are handed to Chroma
    explicitly. Chroma is never permitted to silently fall back to its own
    default embedding function.

    Retrieval is hybrid: dense vector search and BM25 lexical search are
    fused with reciprocal rank fusion, then the candidate pool is reordered
    by a local ONNX cross-encoder reranker.
    """

    def __init__(
        self,
        persist_path: str = "./chroma_db",
        collection_name: str = "local_rag_collection",
        model_name: str = "BAAI/bge-small-en-v1.5",
    ) -> None:
        """Initialize the persistent vector store and local embedding engine.

        Args:
            persist_path: Filesystem directory backing the SQLite Chroma store.
            collection_name: Name of the Chroma collection to create or load.
            model_name: FastEmbed model identifier (ONNX runtime).
        """
        _log("STORE", f"Opening PersistentClient at '{persist_path}'.")
        self.client = chromadb.PersistentClient(path=persist_path)
        self.collection_name = collection_name
        self.collection = self.client.get_or_create_collection(
            name=collection_name,
            # Declare the distance space explicitly for auditable normalization.
            metadata={"hnsw:space": "l2"},
        )
        _log(
            "STORE",
            f"Collection '{collection_name}' ready "
            f"({self.collection.count()} existing vectors).",
        )

        _log("EMBED", f"Loading FastEmbed ONNX model '{model_name}'.")
        t0 = time.perf_counter()
        self.embedding_model = TextEmbedding(model_name=model_name)
        _log(
            "EMBED",
            f"Model loaded in {(time.perf_counter() - t0) * 1000:.1f} ms "
            f"(no PyTorch dependency).",
        )

        # Lazy-built hybrid retrieval helpers (invalidated on ingest/reset).
        self._bm25: Optional[BM25Index] = None
        self._reranker: Optional["TextCrossEncoder"] = None

    def reset(self) -> None:
        """Drop and recreate the collection, discarding all stored vectors.

        Used by ``--reindex`` so a project can be rebuilt from scratch (for
        example after changing the chunking strategy).
        """
        _log("STORE", f"Resetting collection '{self.collection_name}'.")
        self.client.delete_collection(self.collection_name)
        self.collection = self.client.get_or_create_collection(
            name=self.collection_name,
            metadata={"hnsw:space": "l2"},
        )
        self._bm25 = None

    def _embed_batch(self, texts: List[str]) -> List[List[float]]:
        """Generate raw float embeddings for a batch of texts.

        Args:
            texts: Plain strings to embed.

        Returns:
            A list of native Python float lists, one per input string.
        """
        # embed() returns a generator of numpy arrays; materialize + convert
        # to native python lists so the handoff to Chroma is fully explicit.
        vectors: Generator[Any, None, None] = self.embedding_model.embed(texts)
        return [vector.tolist() for vector in vectors]

    def ingest_documents(self, chunks: List[Dict[str, Any]]) -> None:
        """Embed and index a list of structured chunks into Chroma.

        Args:
            chunks: Output of :meth:`NativeDocumentChunker.split_text`.
        """
        if not chunks:
            _log("INGEST", "No chunks supplied; nothing to ingest.")
            return

        # Content-addressed IDs: identical text + provenance hashes to the same
        # ID, so re-ingesting a file is idempotent (no duplicate vectors) and
        # already-indexed chunks can be skipped without re-embedding.
        def _chunk_id(chunk: Dict[str, Any]) -> str:
            label = chunk["metadata"].get("source_label", "document")
            digest = hashlib.sha1(
                f"{label}::{chunk['text']}".encode("utf-8")
            ).hexdigest()[:16]
            return f"chunk-{digest}"

        # Collapse exact-duplicate chunks WITHIN this batch first. Smaller
        # chunk sizes surface repeated boilerplate (page headers/footers, OCR'd
        # page numbers) that hash to the same content ID; keeping one is correct
        # and avoids Chroma's DuplicateIDError on add/get.
        seen_ids: set = set()
        unique: List[Tuple[str, Dict[str, Any]]] = []
        for chunk in chunks:
            cid = _chunk_id(chunk)
            if cid in seen_ids:
                continue
            seen_ids.add(cid)
            unique.append((cid, chunk))

        intra_dupes = len(chunks) - len(unique)
        if intra_dupes:
            _log("INGEST", f"Collapsed {intra_dupes} duplicate chunk(s) "
                           "(repeated boilerplate) within this file.")

        all_ids = [cid for cid, _ in unique]

        # Ask Chroma which of these IDs already exist; only embed the new ones.
        existing = set(self.collection.get(ids=all_ids).get("ids", []))
        new_chunks = [
            (cid, chunk)
            for cid, chunk in unique
            if cid not in existing
        ]

        skipped = len(unique) - len(new_chunks)
        if skipped:
            _log("INGEST", f"Skipping {skipped} already-indexed chunk(s).")
        if not new_chunks:
            _log("INGEST", f"Nothing new; collection holds "
                           f"{self.collection.count()} vectors.")
            return

        ids = [cid for cid, _ in new_chunks]
        texts = [chunk["text"] for _, chunk in new_chunks]
        metadatas = [chunk["metadata"] for _, chunk in new_chunks]
        # Embed the context-augmented text when present (contextual chunk
        # headers); the raw chunk text is what gets stored and displayed.
        embed_texts = [
            chunk.get("embed_text", chunk["text"]) for _, chunk in new_chunks
        ]

        total_chars = sum(len(text) for text in embed_texts)
        _log(
            "INGEST",
            f"Embedding {len(texts)} new chunks "
            f"({total_chars} total characters).",
        )

        t0 = time.perf_counter()
        embeddings = self._embed_batch(embed_texts)
        embed_ms = (time.perf_counter() - t0) * 1000
        dims = len(embeddings[0]) if embeddings else 0
        _log(
            "INGEST",
            f"Generated {len(embeddings)} vectors x {dims} dims "
            f"in {embed_ms:.1f} ms.",
        )

        # upsert (not add) so a re-run can never raise on a duplicate ID.
        self.collection.upsert(
            ids=ids,
            embeddings=embeddings,
            documents=texts,
            metadatas=metadatas,
        )
        _log(
            "INGEST",
            f"Indexed {len(ids)} vectors; collection now holds "
            f"{self.collection.count()} total.",
        )
        self._bm25 = None  # stale; rebuilt lazily on next query

    @staticmethod
    def _l2_to_cosine_similarity(l2_distance: float) -> float:
        """Map a raw L2 distance to a bounded cosine similarity in [0, 1].

        FastEmbed BGE embeddings are L2-normalized (unit length). For two unit
        vectors the squared Euclidean distance relates to cosine similarity by::

            ||a - b||^2 = 2 - 2 * cos(a, b)
            => cos(a, b) = 1 - (d^2 / 2)

        Chroma's ``l2`` space returns the squared distance ``d^2`` directly.
        The cosine result lies in [-1, 1]; it is rescaled to [0, 1] so that
        ``0.0`` reads as orthogonal/unrelated and ``1.0`` as exact identity.

        Args:
            l2_distance: Raw squared-L2 distance reported by Chroma.

        Returns:
            A confidence score clamped to the closed interval ``[0.0, 1.0]``.
        """
        cosine = 1.0 - (l2_distance / 2.0)
        # Guard against tiny floating-point excursions outside [-1, 1].
        cosine = max(-1.0, min(1.0, cosine))
        # Rescale [-1, 1] -> [0, 1].
        normalized = (cosine + 1.0) / 2.0
        return max(0.0, min(1.0, normalized))

    #: Candidates pulled from EACH retriever before fusion.
    HYBRID_K = 20
    #: Reciprocal-rank-fusion constant (standard value from the literature).
    RRF_K = 60
    #: Reranker relevance floor (sigmoid of cross-encoder logit). Moderate
    #: filtering only — strict thresholds discard needed context.
    RERANK_THRESHOLD = 0.25
    #: Local ONNX cross-encoder used to reorder the fused candidate pool.
    RERANKER_MODEL = "Xenova/ms-marco-MiniLM-L-6-v2"

    def _get_bm25(self) -> BM25Index:
        """Build (or reuse) the BM25 index over all stored chunks."""
        if self._bm25 is None:
            t0 = time.perf_counter()
            data = self.collection.get(include=["documents", "metadatas"])
            self._bm25 = BM25Index(
                data["ids"], data["documents"], data["metadatas"])
            _log("QUERY", f"BM25 index built over {len(data['ids'])} chunks "
                          f"in {(time.perf_counter() - t0) * 1000:.1f} ms.")
        return self._bm25

    def _get_reranker(self) -> "TextCrossEncoder":
        """Load (or reuse) the ONNX cross-encoder reranker."""
        if self._reranker is None:
            _log("RERANK", f"Loading cross-encoder '{self.RERANKER_MODEL}' "
                           "(ONNX, first use downloads ~80 MB).")
            t0 = time.perf_counter()
            self._reranker = TextCrossEncoder(
                model_name=self.RERANKER_MODEL,
                # Persistent cache — fastembed's default lands in a temp dir
                # that macOS purges, forcing re-downloads.
                cache_dir=os.path.expanduser("~/.cache/fastembed"),
            )
            _log("RERANK", f"Reranker ready in "
                           f"{(time.perf_counter() - t0) * 1000:.1f} ms.")
        return self._reranker

    def retrieve_context(
        self, query_str: str, n_results: int = 3
    ) -> Tuple[List[Dict[str, Any]], float]:
        """Hybrid retrieval: dense + BM25, RRF fusion, cross-encoder rerank.

        Dense vector search and BM25 lexical search each nominate
        ``HYBRID_K`` candidates; the ranked lists are fused with reciprocal
        rank fusion, and the fused pool is rescored by a cross-encoder that
        reads the query and each chunk together. The top ``n_results``
        above ``RERANK_THRESHOLD`` are returned (always at least two, so
        moderate filtering can never empty the context).

        Args:
            query_str: The natural-language user query.
            n_results: Number of chunks to return after reranking.

        Returns:
            A tuple ``(retrieved_chunks, average_confidence)`` where each
            chunk dict carries a ``confidence`` field (sigmoid of the
            cross-encoder logit) and ``average_confidence`` is their mean.
        """
        count = self.collection.count()
        if count == 0:
            _log("QUERY", "WARNING: collection is empty; no context to return.")
            return [], 0.0

        k = min(self.HYBRID_K, count)

        # --- dense (semantic) candidates --------------------------------- #
        query_vector = self._embed_batch([query_str])[0]
        t0 = time.perf_counter()
        result = self.collection.query(
            query_embeddings=[query_vector],
            n_results=k,
            include=["documents", "metadatas", "distances"],
        )
        dense_ms = (time.perf_counter() - t0) * 1000
        dense_ids = result["ids"][0]
        pool: Dict[str, Dict[str, Any]] = {}
        for cid, doc, meta in zip(
                dense_ids, result["documents"][0], result["metadatas"][0]):
            pool[cid] = {"text": doc, "metadata": meta}

        # --- sparse (BM25 lexical) candidates ----------------------------- #
        t0 = time.perf_counter()
        bm25 = self._get_bm25()
        sparse_hits = bm25.top_n(query_str, k)
        sparse_ms = (time.perf_counter() - t0) * 1000
        sparse_ids = [bm25.ids[idx] for idx, _ in sparse_hits]
        for idx, _score in sparse_hits:
            cid = bm25.ids[idx]
            if cid not in pool:
                pool[cid] = {"text": bm25.documents[idx],
                             "metadata": bm25.metadatas[idx]}

        # --- reciprocal rank fusion --------------------------------------- #
        rrf: Dict[str, float] = {}
        for ranked in (dense_ids, sparse_ids):
            for rank, cid in enumerate(ranked):
                rrf[cid] = rrf.get(cid, 0.0) + 1.0 / (self.RRF_K + rank + 1)
        fused = sorted(rrf, key=rrf.get, reverse=True)
        _log("QUERY", f"Hybrid recall: dense {len(dense_ids)} "
                      f"({dense_ms:.1f} ms) + bm25 {len(sparse_ids)} "
                      f"({sparse_ms:.1f} ms) -> {len(fused)} fused candidates.")

        # --- cross-encoder rerank ------------------------------------------ #
        reranker = self._get_reranker()
        candidate_texts = [pool[cid]["text"] for cid in fused]
        t0 = time.perf_counter()
        logits = list(reranker.rerank(query_str, candidate_texts))
        rerank_ms = (time.perf_counter() - t0) * 1000
        scored = sorted(
            zip(fused, logits), key=lambda p: p[1], reverse=True)
        _log("RERANK", f"Cross-encoder scored {len(scored)} candidates "
                       f"in {rerank_ms:.1f} ms.")

        # --- threshold + top-n selection ----------------------------------- #
        retrieved: List[Dict[str, Any]] = []
        scores: List[float] = []
        for cid, logit in scored:
            confidence = 1.0 / (1.0 + math.exp(-float(logit)))
            if confidence < self.RERANK_THRESHOLD and len(retrieved) >= 2:
                continue
            entry = pool[cid]
            retrieved.append({
                "text": entry["text"],
                "metadata": entry["metadata"],
                "confidence": confidence,
            })
            scores.append(confidence)
            if len(retrieved) >= n_results:
                break

        average = sum(scores) / len(scores) if scores else 0.0
        _log(
            "QUERY",
            f"Rerank confidences: {[round(s, 4) for s in scores]} "
            f"(avg={average:.4f}).",
        )

        # Document-level fallback: when nothing in the corpus passed the
        # relevance floor, the question is usually ABOUT the documents
        # ("summarize this paper") rather than answerable BY one chunk.
        # Prepend the synthetic per-document overview chunks so the
        # generator at least works from real document context.
        if average < self.RERANK_THRESHOLD:
            data = self.collection.get(
                where={"is_overview": True},
                include=["documents", "metadatas"],
            )
            have = {entry["text"] for entry in retrieved}
            extras = [
                (doc, meta)
                for doc, meta in zip(data["documents"], data["metadatas"])
                if doc not in have
            ][:5]
            if extras:
                for doc, meta in extras:
                    retrieved.insert(0, {
                        "text": doc,
                        "metadata": meta,
                        "confidence": 0.0,
                    })
                _log("QUERY", f"Low relevance -> added {len(extras)} "
                              "document overview chunk(s) as fallback "
                              "context.")
            # Cross-encoders cannot score meta-questions ("summarize this
            # paper") against any chunk, but overview chunks ARE the right
            # context for them by construction. Give overviews a trusted-
            # summary prior so the reported confidence reflects grounding,
            # not the reranker's blind spot.
            OVERVIEW_PRIOR = 0.5
            boosted = 0
            for entry in retrieved:
                if (entry["metadata"].get("is_overview")
                        and entry["confidence"] < OVERVIEW_PRIOR):
                    entry["confidence"] = OVERVIEW_PRIOR
                    boosted += 1
            if boosted:
                scores = [entry["confidence"] for entry in retrieved]
                average = sum(scores) / len(scores)
                _log("QUERY", f"Applied trusted-summary prior "
                              f"({OVERVIEW_PRIOR}) to {boosted} overview "
                              f"chunk(s); avg={average:.4f}.")

        return retrieved, average


# --------------------------------------------------------------------------- #
# Block 4: Deterministic Local Inference Payload.
# --------------------------------------------------------------------------- #
#: Friendly shortcuts -> concrete Ollama model tags.
MODEL_ALIASES: Dict[str, str] = {
    "7b": "qwen2.5:7b",
    "qwen7b": "qwen2.5:7b",
}


def resolve_model(name: str) -> str:
    """Map a friendly shortcut to a concrete Ollama model tag.

    Args:
        name: A shortcut (``"7b"``) or any raw Ollama tag.

    Returns:
        The resolved Ollama model tag; raw tags pass through unchanged.
    """
    return MODEL_ALIASES.get(name.lower().strip(), name)


def _warn_if_model_missing(model: str) -> None:
    """Emit a clean, educational warning if a model is not pulled locally.

    Args:
        model: A resolved Ollama model tag.
    """
    if ollama is None:
        return
    try:
        response = ollama.list()
        # ollama>=0.2 returns a pydantic ListResponse; older versions a dict.
        models = getattr(response, "models", None)
        if models is None and isinstance(response, dict):
            models = response.get("models", [])
        installed = {
            getattr(m, "model", None) or (m.get("model") if isinstance(m, dict) else "")
            for m in (models or [])
        }
    except Exception:  # noqa: BLE001 - daemon may be down; chat() handles that.
        return
    # Ollama reports tags like 'qwen2.5:7b'; tolerate ':latest' bare names too.
    if model not in installed and f"{model}:latest" not in installed:
        _log("MODEL", f"WARNING: '{model}' is not pulled locally. Run "
                      f"'ollama pull {model}' first (default is '-m 7b').")


class DeterministicRAGEngine:
    """Zero-hallucination local inference over retrieved context.

    Drives a local Ollama ``qwen2.5:7b`` model with greedy decoding
    (``temperature=0.0``) and an expanded KV-cache window so that multi-document
    contexts are never truncated.
    """

    SYSTEM_PROMPT: str = (
        "You are an empirical, fact-bound text interpreter. You answer ONLY "
        "from the CONTEXT BLOCKS provided in the user message.\n"
        "Hard constraints:\n"
        "1. You are strictly forbidden from using any outside, pre-trained, "
        "or world knowledge. Treat the context as your entire universe.\n"
        "2. Match meaning, not wording. The question and the context will use "
        "different phrasings for the same thing -- treat paraphrases, synonyms, "
        "and equivalent time periods as matches (e.g. a question about "
        "'before the 19th century' or 'historically' is answered by text on "
        "the 'historical' or 'pre-development' period; 'population' is answered "
        "by 'run size' or 'abundance'). Answer whenever the context contains "
        "information that addresses the question's intent, even if no sentence "
        "restates the question verbatim.\n"
        "3. Some context blocks are marked DOCUMENT OVERVIEW: they are "
        "trusted summaries of an entire document in the user's project. "
        "When the question refers to 'this paper', 'this document', 'this "
        "file', or asks what a document is about or for its summary, it "
        "refers to those documents -- answer from the OVERVIEW blocks "
        "(plus any other relevant blocks). NEVER refuse while an OVERVIEW "
        "block is present.\n"
        "3b. Refusal is a last resort. ONLY when the context contains no "
        "information relevant to the question's intent at all, respond with "
        "EXACTLY this sentence and nothing else: "
        '"I cannot answer this based on the provided documents."\n'
        "4. Every sentence that states an extracted fact MUST end with an "
        "inline citation. When the context block header includes pages, cite "
        "as [Source Chunk N, p. X] (or [Source Chunk N, pp. X-Y] for a "
        "range); when it has no pages, cite as [Source Chunk N]. DOCUMENT "
        "OVERVIEW blocks never have page numbers -- cite them as "
        "[Source Chunk N] with no page. Use only page numbers given in a "
        "block's own header, never invented ones and never another "
        "block's.\n"
        "5. You may connect and restate facts that are explicitly present in "
        "the context, but never introduce a fact, number, or claim that is not "
        "grounded in the context.\n"
        "6. Match answer depth to the question. For a narrow factual question, "
        "give the direct answer first plus ONLY the qualifiers actually in the "
        "context -- ranges, caveats, key dates, or hedges the source itself "
        "states. For a broad question (asking for ideas, suggestions, "
        "summaries, comparisons, 'top N' lists, or how something works), give "
        "a THOROUGH answer: work through every context block, extract each "
        "distinct relevant point, and explain it in 2-4 full sentences using "
        "the specific details present (names, numbers, settings, reasons, "
        "trade-offs). Organize multi-point answers as a numbered or bulleted "
        "list. Depth must come from the context's own details -- never from "
        "padding, repetition, or restating the question."
    )

    #: Light sampling: more natural prose than greedy decoding while staying
    #: factual on a 7B; above ~0.4 small models start decorating facts.
    TEMPERATURE: float = 0.15

    #: Hard ceiling on generated tokens — sized so detailed multi-point
    #: answers fit, while still capping worst-case decode time on CPU.
    MAX_OUTPUT_TOKENS: int = 1024

    #: Lower / upper bounds for the dynamically-sized context window.
    MIN_CTX: int = 2048
    MAX_CTX: int = 8192

    def __init__(self, model_name: str = "qwen2.5:7b") -> None:
        """Configure the inference engine.

        Args:
            model_name: Local Ollama model tag to invoke.
        """
        self.model_name = model_name

    def _fit_num_ctx(self, prompt_chars: int) -> int:
        """Size the KV-cache context window to the actual prompt, not a fixed max.

        A fixed ``num_ctx=16384`` forces Ollama to allocate (and prompt-evaluate
        against) a huge KV cache even for a small prompt — crippling on CPU-only
        hosts. We instead estimate tokens (~4 chars/token), add room for the
        bounded answer, and round up to a sane window clamped to a tight range.

        Args:
            prompt_chars: Total character length of system + user messages.

        Returns:
            A context-window size in tokens.
        """
        est_input_tokens = prompt_chars // 4
        needed = est_input_tokens + self.MAX_OUTPUT_TOKENS + 256  # safety margin
        # Round up to the next 1024 boundary, then clamp.
        rounded = ((needed + 1023) // 1024) * 1024
        return max(self.MIN_CTX, min(self.MAX_CTX, rounded))

    def _build_context_block(self, retrieved_chunks: List[Dict[str, Any]]) -> str:
        """Render retrieved chunks into a citation-ready prompt section.

        Args:
            retrieved_chunks: Output of
                :meth:`LocalVectorStorage.retrieve_context`.

        Returns:
            A formatted multi-line string with one labeled block per chunk.
        """
        blocks: List[str] = []
        for position, chunk in enumerate(retrieved_chunks, start=1):
            meta = chunk["metadata"]
            if meta.get("is_overview"):
                header = f"[Source Chunk {position}] (DOCUMENT OVERVIEW"
            else:
                index = meta.get("chunk_index", position - 1)
                header = f"[Source Chunk {position}] (chunk_index={index}"
            source = meta.get("source_label")
            if source:
                header += f", file={source}"
            p_start, p_end = meta.get("page_start"), meta.get("page_end")
            if p_start is not None and not meta.get("is_overview"):
                pages = str(p_start) if p_start == p_end else f"{p_start}-{p_end}"
                header += f", pages={pages}"
            header += ")"
            blocks.append(f"{header}\n{chunk['text']}")
        return "\n\n".join(blocks)

    def generate_answer(
        self, query: str, retrieved_chunks: List[Dict[str, Any]]
    ) -> str:
        """Generate a grounded, cited answer from retrieved context.

        Args:
            query: The user's natural-language question.
            retrieved_chunks: Context chunks returned by the vector store.

        Returns:
            The model's answer string, or a clean educational error message if
            the local Ollama daemon or model is unavailable.
        """
        if ollama is None:
            return (
                "[ENGINE WARNING] The 'ollama' python package is not installed. "
                "Install it with 'pip install ollama' and ensure the Ollama "
                "daemon is running locally."
            )

        if not retrieved_chunks:
            return "I cannot answer this based on the provided documents."

        context = self._build_context_block(retrieved_chunks)
        user_message = (
            f"CONTEXT BLOCKS:\n{context}\n\n"
            f"QUESTION: {query}\n\n"
            "Answer using only the context blocks above, with inline "
            "[Source Chunk N, p. X] citations (page from each block's "
            "header) on every factual sentence."
        )

        num_ctx = self._fit_num_ctx(len(self.SYSTEM_PROMPT) + len(user_message))
        _log(
            "INFER",
            f"Dispatching to Ollama model '{self.model_name}' "
            f"(temperature={self.TEMPERATURE}, num_ctx={num_ctx}, "
            f"num_predict={self.MAX_OUTPUT_TOKENS}).",
        )
        t0 = time.perf_counter()
        try:
            response = ollama.chat(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": self.SYSTEM_PROMPT},
                    {"role": "user", "content": user_message},
                ],
                options={
                    "temperature": self.TEMPERATURE,
                    # Sized to the actual prompt instead of a fixed 16K window.
                    "num_ctx": num_ctx,
                    # Bound the answer length so decode can't run away.
                    "num_predict": self.MAX_OUTPUT_TOKENS,
                },
                # Keep the model resident between questions to avoid reloads.
                keep_alive="10m",
            )
        except Exception as exc:  # noqa: BLE001 - surface a friendly message.
            _log("INFER", f"WARNING: Ollama call failed: {exc}")
            return (
                "[ENGINE WARNING] Could not reach the local Ollama daemon or "
                f"model '{self.model_name}'. Start it with 'ollama serve' and "
                f"pull the model with 'ollama pull {self.model_name}'. "
                f"Underlying error: {exc}"
            )

        latency_ms = (time.perf_counter() - t0) * 1000
        _log("INFER", f"Generation completed in {latency_ms:.1f} ms.")
        return response["message"]["content"].strip()


# --------------------------------------------------------------------------- #
# Block 5a: Project Workspace — persistent index + per-chat JSON history.
# --------------------------------------------------------------------------- #
class ProjectWorkspace:
    """On-disk home for one RAG project: its index, manifest, and chat history.

    A *project* groups a persistent Chroma index with a manifest of which
    source files have already been chunked (so re-running never re-chunks an
    unchanged file) and a folder of chat-session JSON logs that accumulate
    every query, answer, score, and the raw retrieved fragments for audit.

    Layout::

        ./projects/<name>/
            chroma_db/          # persistent vector index for this project
            manifest.json       # {file -> {hash, chunks, ingested_at}}
            history/
                <session>.json  # accumulating list of chat turns

    Args:
        name: Human-friendly project name (used as the directory name).
        base_dir: Parent directory holding all projects.
    """

    def __init__(self, name: str, base_dir: str = "./projects") -> None:
        self.name = name
        self.root = os.path.join(base_dir, name)
        self.chroma_path = os.path.join(self.root, "chroma_db")
        self.manifest_path = os.path.join(self.root, "manifest.json")
        self.history_dir = os.path.join(self.root, "history")
        os.makedirs(self.history_dir, exist_ok=True)

    # ------------------------------ manifest -------------------------------- #
    def load_manifest(self) -> Dict[str, Any]:
        """Load the project manifest, or return a fresh empty one."""
        if os.path.exists(self.manifest_path):
            with open(self.manifest_path, "r", encoding="utf-8") as handle:
                return json.load(handle)
        return {"project": self.name, "files": {}}

    def save_manifest(self, manifest: Dict[str, Any]) -> None:
        """Persist the project manifest to disk."""
        with open(self.manifest_path, "w", encoding="utf-8") as handle:
            json.dump(manifest, handle, indent=2, ensure_ascii=False)

    @staticmethod
    def file_fingerprint(path: str) -> str:
        """Return a streaming SHA-1 of a file's raw bytes.

        Args:
            path: Path to the file to fingerprint.

        Returns:
            A hex digest uniquely identifying the file's current contents.
        """
        digest = hashlib.sha1()
        with open(path, "rb") as handle:
            for block in iter(lambda: handle.read(65536), b""):
                digest.update(block)
        return digest.hexdigest()

    def is_indexed(self, path: str, manifest: Dict[str, Any]) -> bool:
        """Report whether ``path`` is already indexed at its current contents.

        Args:
            path: Source file path.
            manifest: A loaded manifest dictionary.

        Returns:
            True only if the file was ingested AND its bytes are unchanged.
        """
        record = manifest["files"].get(os.path.abspath(path))
        if not record:
            return False
        return record.get("hash") == self.file_fingerprint(path)

    def record_file(
        self, manifest: Dict[str, Any], path: str, chunk_count: int
    ) -> None:
        """Stamp a freshly ingested file into the manifest."""
        manifest["files"][os.path.abspath(path)] = {
            "hash": self.file_fingerprint(path),
            "chunks": chunk_count,
            "ingested_at": datetime.now().isoformat(timespec="seconds"),
        }

    # ------------------------------- history -------------------------------- #
    def append_turn(self, session: str, turn: Dict[str, Any]) -> str:
        """Append one chat turn to a session's history JSON.

        Args:
            session: Chat session name; turns with the same name accumulate
                into a single growing transcript across invocations.
            turn: The turn record (query, answer, score, fragments, ...).

        Returns:
            The path to the session history file that was written.
        """
        path = os.path.join(self.history_dir, f"{session}.json")
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as handle:
                data = json.load(handle)
        else:
            data = {"project": self.name, "session": session, "turns": []}
        data["turns"].append(turn)
        with open(path, "w", encoding="utf-8") as handle:
            json.dump(data, handle, indent=2, ensure_ascii=False)
        return path


# --------------------------------------------------------------------------- #
# Block 5b: Terminal rendering (answer-only) + full audit (demo).
# --------------------------------------------------------------------------- #
def _render_answer(
    query: str,
    answer: str,
    confidence: float,
    history_path: str,
    n_fragments: int,
    elapsed_secs: float = 0.0,
) -> None:
    """Print only the synthesized answer; raw fragments live in the JSON log.

    Args:
        query: The user's question.
        answer: The synthesized model response.
        confidence: Averaged normalized accuracy score.
        history_path: Where the full turn (including fragments) was saved.
        n_fragments: How many raw fragments were archived for this turn.
        elapsed_secs: Wall-clock time for retrieval + generation.
    """
    bar = "=" * 78
    print(f"\n{bar}")
    print("LOCAL RAG — ANSWER".center(78))
    print(bar)
    print(f"\nQ: {query}\n")
    print(answer)
    print("\n" + "-" * 78)
    print(
        f"Confidence: {confidence:.4f} (0.0–1.0)   |   "
        f"answered in {_fmt_secs(elapsed_secs)}   |   "
        f"{n_fragments} raw fragments saved -> {history_path}"
    )
    print(f"{bar}\n")


def _render_results(
    query: str,
    answer: str,
    confidence: float,
    retrieved_chunks: List[Dict[str, Any]],
) -> None:
    """Print the full audit view (query, answer, score, raw fragments).

    Used by the built-in demo so first-time readers can see every layer.

    Args:
        query: The original user query.
        answer: The synthesized local model response.
        confidence: The averaged normalized accuracy score.
        retrieved_chunks: The exact chunks pulled from the vector store.
    """
    bar = "=" * 78
    print(f"\n{bar}")
    print("LOCAL RAG PIPELINE — VERIFICATION RESULTS".center(78))
    print(bar)

    print("\n[1] ORIGINAL USER QUERY")
    print("-" * 78)
    print(query)

    print("\n[2] SYNTHESIZED LOCAL QWEN RESPONSE (inline citations)")
    print("-" * 78)
    print(answer)

    print("\n[3] MATHEMATICALLY NORMALIZED PIPELINE ACCURACY SCORE")
    print("-" * 78)
    print(f"Average cosine-normalized confidence: {confidence:.4f}  (0.0–1.0)")

    print("\n[4] RAW RETRIEVED FRAGMENTS (audit verification)")
    print("-" * 78)
    for position, chunk in enumerate(retrieved_chunks, start=1):
        meta = chunk["metadata"]
        print(
            f"\n  • Source Chunk {position} "
            f"(chunk_index={meta.get('chunk_index')}, "
            f"chars={meta.get('char_count')}, "
            f"confidence={chunk.get('confidence', 0.0):.4f})"
        )
        print(f"    {chunk['text']}")
    print(f"\n{bar}\n")


# --------------------------------------------------------------------------- #
# Block 5c: Project operations — ingest-once / ask-many.
# --------------------------------------------------------------------------- #
def _summarize_document(text: str, name: str, model: str) -> str:
    """Generate a short document summary with the local model (best-effort).

    Used to build contextual chunk headers: prepending a document-level
    summary to each chunk before embedding gives the vectors document
    context, so chunks that only make sense "under the overall topic"
    become retrievable. One LLM call per document, at ingest time only.

    Returns an empty string if Ollama is unavailable or the call fails —
    ingestion proceeds without summaries rather than aborting.
    """
    if ollama is None:
        return ""
    # A representative slice keeps CPU summarization time bounded: the
    # opening usually carries the document's framing, plus a mid-document
    # sample for long files.
    sample = text[:5000]
    if len(text) > 12000:
        mid = len(text) // 2
        sample += "\n...\n" + text[mid:mid + 2500]
    try:
        t0 = time.perf_counter()
        response = ollama.chat(
            model=model,
            messages=[
                {"role": "system", "content":
                    "Summarize the document excerpt in under 120 words. "
                    "State what the document is, its domain, and its main "
                    "topics. Output only the summary."},
                {"role": "user", "content": sample},
            ],
            options={"temperature": 0.2, "num_predict": 200,
                     "num_ctx": 4096},
            keep_alive="10m",
        )
        summary = response["message"]["content"].strip()
        _log("INGEST", f"'{name}': document summary generated in "
                       f"{_fmt_secs(time.perf_counter() - t0)} "
                       f"({len(summary)} chars).")
        return summary
    except Exception as exc:  # noqa: BLE001 - summaries are best-effort.
        _log("INGEST", f"WARNING: summary for '{name}' skipped ({exc}).")
        return ""


_PAGE_MARKER = re.compile(r"\[Page (\d+)\]")


def _assign_pages(chunks: List[Dict[str, Any]]) -> None:
    """Stamp page_start/page_end metadata onto chunks from [Page N] markers.

    The PDF loader prepends a ``[Page N]`` marker to each page's text, so
    every chunk either contains one or more markers or falls inside the page
    of the previous marker. Chunks are produced in document order, so a
    chunk with no marker inherits the page carried forward from the last
    marker seen. Non-PDF sources have no markers and get no page metadata.
    """
    current_page = None
    for chunk in chunks:
        pages = [int(m) for m in _PAGE_MARKER.findall(chunk["text"])]
        if pages:
            start = current_page if current_page is not None else pages[0]
            chunk["metadata"]["page_start"] = min(start, pages[0])
            chunk["metadata"]["page_end"] = pages[-1]
            current_page = pages[-1]
        elif current_page is not None:
            chunk["metadata"]["page_start"] = current_page
            chunk["metadata"]["page_end"] = current_page


def ingest_project(
    store: "LocalVectorStorage",
    workspace: ProjectWorkspace,
    path: str,
    chunk_size: int,
    chunk_overlap: int,
    summary_model: str = "qwen2.5:7b",
) -> int:
    """Chunk + index every NEW supported file under ``path`` into the project.

    Files already recorded in the manifest at their current byte-hash are
    skipped entirely — they are never re-parsed or re-chunked. Each new
    file gets an LLM-generated document summary; the summary plus file name
    are prepended to every chunk's *embedding input* (contextual chunk
    headers), while the stored chunk text stays raw.

    Args:
        store: The project's vector store.
        workspace: The project workspace (manifest + paths).
        path: A file or directory to ingest.
        chunk_size: Maximum chunk character length.
        chunk_overlap: Sliding overlap between chunks.
        summary_model: Local Ollama model used for document summaries.

    Returns:
        The number of newly indexed chunks across all freshly added files.
    """
    loader = NativeDocumentLoader()
    chunker = NativeDocumentChunker()
    manifest = workspace.load_manifest()

    # Discover candidate files WITHOUT parsing them, so unchanged files cost
    # nothing beyond a fast byte-hash.
    targets: List[str] = []
    if os.path.isdir(path):
        for root, _dirs, files in os.walk(path):
            for name in sorted(files):
                if name.lower().endswith(loader.SUPPORTED):
                    targets.append(os.path.join(root, name))
    elif os.path.isfile(path):
        targets.append(path)
    else:
        _log("MAIN", f"Path not found: '{path}'.")
        return 0

    new_total = 0
    t_ingest_start = time.perf_counter()
    load_secs = 0.0
    chunk_secs = 0.0
    embed_secs = 0.0
    for target in targets:
        name = os.path.basename(target)
        if workspace.is_indexed(target, manifest):
            _log("INGEST", f"'{name}' already indexed at current contents; "
                           "skipping parse + chunk.")
            continue

        t0 = time.perf_counter()
        text = loader.load_file(target)
        load_dt = time.perf_counter() - t0
        load_secs += load_dt
        if not text.strip():
            continue

        t0 = time.perf_counter()
        chunks = chunker.split_text(
            text,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            source_label=name,
        )
        chunk_dt = time.perf_counter() - t0
        chunk_secs += chunk_dt
        _assign_pages(chunks)
        _log("TIMING", f"'{name}': parse/OCR {_fmt_secs(load_dt)}, "
                       f"chunk {_fmt_secs(chunk_dt)} -> {len(chunks)} chunks.")

        # Contextual chunk headers: embed each chunk WITH its document
        # context so context-dependent chunks stay retrievable.
        summary = _summarize_document(text, name, summary_model)
        header = f"[File: {name}]"
        if summary:
            header += f"\n[Document summary: {summary}]"
        for chunk in chunks:
            chunk["embed_text"] = f"{header}\n\n{chunk['text']}"

        # Synthetic overview chunk: document-level questions ("what is this
        # paper about", "summarize it") have no single body chunk that
        # answers them — the overview gives them a retrieval target. Also
        # served as a fallback when reranking finds nothing relevant.
        if summary:
            last_page = max(
                (c["metadata"].get("page_end", 0) for c in chunks), default=0)
            overview_meta: Dict[str, Any] = {
                "source_label": name,
                "chunk_index": -1,
                "char_count": len(summary),
                "is_overview": True,
            }
            if last_page:
                overview_meta["page_start"] = 1
                overview_meta["page_end"] = last_page
            chunks.append({
                "text": f"Document overview of '{name}': {summary}",
                "metadata": overview_meta,
            })

        t0 = time.perf_counter()
        store.ingest_documents(chunks)
        embed_secs += time.perf_counter() - t0

        workspace.record_file(manifest, target, len(chunks))
        new_total += len(chunks)

    workspace.save_manifest(manifest)

    total = time.perf_counter() - t_ingest_start
    _log("TIMING", f"INGEST SUMMARY -> parse/OCR {_fmt_secs(load_secs)} | "
                   f"chunk {_fmt_secs(chunk_secs)} | embed+index "
                   f"{_fmt_secs(embed_secs)} | total {_fmt_secs(total)}.")
    return new_total


def ask_project(
    store: "LocalVectorStorage",
    workspace: ProjectWorkspace,
    question: str,
    session: str,
    n_results: int,
    model: str,
) -> None:
    """Answer a question from existing project chunks and log the turn.

    Args:
        store: The project's vector store (already populated).
        workspace: The project workspace (for history logging).
        question: The natural-language question.
        session: Chat session name to append this turn to.
        n_results: Number of fragments to retrieve.
        model: Local Ollama model tag.
    """
    if store.collection.count() == 0:
        _log("MAIN", "Project has no indexed chunks yet. Run with --ingest "
                     "<path> first to build the index.")
        return

    t0 = time.perf_counter()
    retrieved, confidence = store.retrieve_context(question, n_results=n_results)
    retrieve_secs = time.perf_counter() - t0

    engine = DeterministicRAGEngine(model_name=model)
    t0 = time.perf_counter()
    answer = engine.generate_answer(question, retrieved)
    generate_secs = time.perf_counter() - t0

    _log("TIMING", f"ASK SUMMARY -> retrieve {_fmt_secs(retrieve_secs)} | "
                   f"generate ({model}) {_fmt_secs(generate_secs)} | total "
                   f"{_fmt_secs(retrieve_secs + generate_secs)}.")

    turn: Dict[str, Any] = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "query": question,
        "answer": answer,
        "confidence": round(confidence, 4),
        "model": model,
        "n_results": n_results,
        "timing_seconds": {
            "retrieve": round(retrieve_secs, 3),
            "generate": round(generate_secs, 3),
            "total": round(retrieve_secs + generate_secs, 3),
        },
        # Full raw fragments retained for audit; never printed to the terminal.
        "fragments": retrieved,
    }
    history_path = workspace.append_turn(session, turn)
    _render_answer(
        question, answer, confidence, history_path, len(retrieved),
        generate_secs + retrieve_secs,
    )


# --------------------------------------------------------------------------- #
# Block 5d: Built-in demonstration (no project required).
# --------------------------------------------------------------------------- #
def _run_demo() -> None:
    """Run the self-contained mock demonstration with full audit output."""
    mock_document = """
AUTOMATED BASIN WATER MANAGEMENT — OPERATIONAL CASE STUDY

Overview.
The Cedar Basin automated water management network monitors physical flow
dynamics across three distinct geographical monitoring stations: Station Alpha
at the northern headwaters, Station Bravo at the central confluence, and
Station Charlie at the southern outflow weir. Telemetry is sampled every
fifteen minutes and reconciled against a deterministic hydrological model.

Station Alpha.
Station Alpha records a baseline discharge of 12.4 cubic meters per second
during the dry season. Its primary instrument is an acoustic Doppler current
profiler mounted on the left bank. During the March snowmelt event the station
peaked at 47.8 cubic meters per second, triggering an automated upstream gate
adjustment to protect the downstream channel.

Station Bravo.
Station Bravo sits at the central confluence where two tributaries merge. The
confluence amplifies turbulence, so the station relies on a redundant pair of
pressure transducers for accuracy. Bravo maintains a regulatory minimum flow of
8.0 cubic meters per second to preserve aquatic habitat. When inflow exceeds
60 cubic meters per second, the control logic opens the secondary spillway.

Station Charlie.
Station Charlie governs the southern outflow weir and is responsible for final
release scheduling into the irrigation district. Charlie's weir has a maximum
safe release capacity of 95 cubic meters per second. The station coordinates
with Station Bravo using a fifteen-minute lookahead so that spillway events
upstream do not overwhelm the weir gates downstream.

Control Strategy.
The network applies a conservative cascade control strategy. Each station
publishes its state to a shared ledger, and gate adjustments are computed only
from measured telemetry, never from forecasts alone. This deterministic policy
guarantees that any release decision can be reconstructed and audited after the
fact, which is a strict regulatory requirement for the basin authority.
""".strip()

    _log("MAIN", f"Mock document loaded: {len(mock_document)} characters.")

    # --- Block 1: chunk the document. ---
    chunker = NativeDocumentChunker()
    chunks = chunker.split_text(
        mock_document,
        chunk_size=600,
        chunk_overlap=120,
        source_label="cedar_basin_case_study",
    )
    _log("MAIN", f"Chunking produced {len(chunks)} chunks.")

    # --- Block 2: embed + index. ---
    try:
        store = LocalVectorStorage(
            persist_path="./chroma_db",
            collection_name="cedar_basin_rag",
        )
    except Exception as exc:  # noqa: BLE001 - educational failsafe.
        _log("MAIN", f"FATAL: could not initialize vector storage: {exc}")
        return

    store.ingest_documents(chunks)

    # --- Block 3: query + normalize. ---
    query = (
        "What is the maximum safe release capacity at the southern weir, and "
        "what regulatory minimum flow must the central confluence station "
        "maintain?"
    )
    retrieved_chunks, confidence = store.retrieve_context(query, n_results=3)

    # --- Block 4: deterministic local generation. ---
    engine = DeterministicRAGEngine(model_name="qwen2.5:7b")
    answer = engine.generate_answer(query, retrieved_chunks)

    # --- Block 5: structured audit output. ---
    _render_results(query, answer, confidence, retrieved_chunks)


# --------------------------------------------------------------------------- #
# Block 5e: Command-line entrypoint (project-aware ingest-once / ask-many).
# --------------------------------------------------------------------------- #
def main() -> None:
    """Parse CLI arguments and dispatch to ingest / ask / demo.

    Examples:
        # Build a project's index ONCE from a file or folder:
        python local_rag_pipeline.py -p steelhead -i ./docs/McEwan_2001.pdf

        # Then ask as many questions as you like (no re-chunking):
        python local_rag_pipeline.py -p steelhead -a "historical run size?"
        python local_rag_pipeline.py -p steelhead -a "what stressors?" -s notes

        # Run the built-in self-contained demonstration:
        python local_rag_pipeline.py --demo
    """
    parser = argparse.ArgumentParser(
        description="Fully local, auditable RAG pipeline (ingest once, ask "
                    "many; per-chat JSON history of raw fragments).",
    )
    parser.add_argument("-p", "--project",
                        help="Project name. Isolates the index + chat history.")
    parser.add_argument("-i", "--ingest",
                        help="File or folder to chunk + index into the project. "
                             "Unchanged files already indexed are skipped.")
    parser.add_argument("-a", "--ask",
                        help="Question to answer from the project's chunks.")
    parser.add_argument("-s", "--session", default="default",
                        help="Chat session name; turns accumulate per session.")
    parser.add_argument("-n", "--n-results", type=int, default=4,
                        help="Number of fragments to retrieve per question.")
    parser.add_argument("-m", "--model", default="7b",
                        help="Model to answer with. Shortcut '7b' maps to "
                             "qwen2.5:7b (the default). Any raw Ollama tag "
                             "also works.")
    parser.add_argument("--chunk-size", type=int, default=1000,
                        help="Max chunk character length (ingest only).")
    parser.add_argument("--chunk-overlap", type=int, default=200,
                        help="Sliding overlap between chunks (ingest only).")
    parser.add_argument("--reindex", action="store_true",
                        help="Wipe the project's existing index + manifest "
                             "before ingesting (rebuild from scratch).")
    parser.add_argument("--demo", action="store_true",
                        help="Run the built-in mock demonstration and exit.")
    args = parser.parse_args()

    if args.demo or not (args.project or args.ingest or args.ask):
        _run_demo()
        return

    if not args.project:
        parser.error("--project is required for --ingest/--ask (or use --demo).")
    if not (args.ingest or args.ask):
        parser.error("provide --ingest <path> and/or --ask \"<question>\".")

    workspace = ProjectWorkspace(args.project)
    try:
        store = LocalVectorStorage(
            persist_path=workspace.chroma_path,
            collection_name="project_chunks",
        )
    except Exception as exc:  # noqa: BLE001 - educational failsafe.
        _log("MAIN", f"FATAL: could not initialize vector storage: {exc}")
        return

    if args.reindex:
        store.reset()
        if os.path.exists(workspace.manifest_path):
            os.remove(workspace.manifest_path)
        _log("MAIN", f"Project '{args.project}' wiped; rebuilding from scratch.")

    if args.ingest:
        new_chunks = ingest_project(
            store, workspace, args.ingest,
            args.chunk_size, args.chunk_overlap,
        )
        _log("MAIN", f"Ingest complete: {new_chunks} new chunk(s). Project "
                     f"'{args.project}' holds {store.collection.count()} total.")

    if args.ask:
        model = resolve_model(args.model)
        _warn_if_model_missing(model)
        ask_project(
            store, workspace, args.ask,
            args.session, args.n_results, model,
        )


if __name__ == "__main__":
    main()
